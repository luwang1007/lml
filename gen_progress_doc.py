from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from datetime import datetime

def set_chinese_font(run, font_name):
    """设置中文字体"""
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

def generate_report():
    doc = Document()

    # ==========================================
    # 1. 全局样式设置 (字体与段落)
    # ==========================================
    # 设置正文默认样式
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Arial'  # 英文数字字体
    normal_style.font.size = Pt(11)
    normal_style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')  # 中文字体
    normal_style.paragraph_format.line_spacing = 1.5  # 1.5倍行距
    normal_style.paragraph_format.space_after = Pt(10)  # 段后间距

    # 商务主色调：深蓝色 (Business Blue)
    PRIMARY_COLOR = RGBColor(31, 73, 125)

    # 标题样式统一设置
    for i in range(1, 4):
        h_style = doc.styles[f'Heading {i}']
        h_style.font.color.rgb = PRIMARY_COLOR
        h_style.font.name = 'Arial'
        h_style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        
    doc.styles['Heading 1'].font.size = Pt(16)
    doc.styles['Heading 1'].paragraph_format.space_before = Pt(24)
    doc.styles['Heading 1'].paragraph_format.space_after = Pt(12)

    # ==========================================
    # 2. 封面设计
    # ==========================================
    doc.add_paragraph('\n' * 4)  # 顶部留白

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run('商贸销售数据分析与智能预测系统')
    title_run.font.size = Pt(24)
    title_run.bold = True
    title_run.font.color.rgb = PRIMARY_COLOR
    set_chinese_font(title_run, '微软雅黑')

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_p.add_run('项 目 进 度 报 告')
    sub_run.font.size = Pt(16)
    sub_run.font.color.rgb = RGBColor(89, 89, 89) # 深灰色
    sub_run.bold = True
    set_chinese_font(sub_run, '微软雅黑')

    doc.add_paragraph('\n' * 8)  # 中间留白

    info_p = doc.add_paragraph()
    info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_run = info_p.add_run(f'汇报人：刘明亮\n生成日期：{datetime.now().strftime("%Y年%m月%d日")}')
    info_run.font.size = Pt(14)
    info_run.font.color.rgb = RGBColor(64, 64, 64)
    set_chinese_font(info_run, '微软雅黑')

    doc.add_page_break() # 换页

    # ==========================================
    # 3. 页眉设置
    # ==========================================
    section = doc.sections[0]
    header = section.header
    header_p = header.paragraphs[0]
    header_p.text = '商贸销售数据分析与智能预测系统 - 内部资料'
    header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_p.style.font.size = Pt(9)
    header_p.style.font.color.rgb = RGBColor(128, 128, 128)

    # ==========================================
    # 4. 正文内容
    # ==========================================
    doc.add_heading('一、总体进度概览', level=1)
    p = doc.add_paragraph()
    run1 = p.add_run('当前完成度：')
    run1.bold = True
    run2 = p.add_run('约 50%')
    run2.font.color.rgb = RGBColor(0, 150, 136) # 进度用商务绿强调
    run2.bold = True
    
    doc.add_paragraph('项目基于 Kaggle Favorita Store Sales 真实零售数据集，采用 Flask + ECharts + ARIMA/Prophet/LSTM 多模型预测架构。后端全部模块已完成并通过代码审查，前端页面已完成并修复关键 Bug，数据集已准备就绪。')

    # ==========================================
    # 5. 表格美化 (各步骤完成情况)
    # ==========================================
    doc.add_heading('二、各步骤完成情况', level=1)
    
    table = doc.add_table(rows=1, cols=4)
    # 使用 Word 内置的专业商务表格样式
    table.style = 'Light Shading Accent 1' # 也可以用 'Grid Table 4 Accent 1'
    
    # 调整列宽 (注意: Word表格列宽受内容影响，这只是建议值)
    table.autofit = False
    table.columns[0].width = Cm(2.0)
    table.columns[1].width = Cm(4.5)
    table.columns[2].width = Cm(2.5)
    table.columns[3].width = Cm(7.0)

    hdr = table.rows[0].cells
    headers = ['Step', '模块名称', '状态', '说明']
    for i, text in enumerate(headers):
        hdr[i].text = text
        hdr[i].paragraphs[0].runs[0].font.bold = True
        hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    steps = [
        ('Step 1', '项目骨架与环境搭建', '✅ 完成', 'config.py、prepare_data.py、目录结构、离线静态资源'),
        ('Step 2', '数据清洗与预处理', '✅ 完成', 'DataProcessor：加载/校验/清洗/特征工程/落盘'),
        ('Step 3', '数据分析与可视化', '✅ 完成', 'DataAnalyzer：P0 全部 + P1 可选图表（含促销效果）'),
        ('Step 4', 'ARIMA 预测模型', '✅ 完成', '自动定阶、网格搜索、滚动评估、序列化'),
        ('Step 5', 'Prophet 预测模型', '✅ 完成', '节假日加载、促销回归量、分量分解'),
        ('Step 6', 'LSTM 预测模型', '✅ 完成', '双层 LSTM、早停、一步预测评估'),
        ('Step 7', '模型评估与对比', '✅ 完成', 'ModelEvaluator：五大指标、雷达图、综合排名'),
        ('Step 8', 'Flask 后端 API', '✅ 完成', '全部接口：上传/分析/预测/进度/结果/导出'),
        ('Step 9', '前端页面与 ECharts', '✅ 完成', '4页模板 + JS 全部实现，修复 6 处数据结构 Bug'),
        ('Step 10', '系统集成测试', '⚠️ 部分', '测试框架已建立，部分测试数据文件待补充'),
    ]

    for s in steps:
        row = table.add_row().cells
        for i, val in enumerate(s):
            row[i].text = val
            row[i].paragraphs[0].paragraph_format.space_before = Pt(6)
            row[i].paragraphs[0].paragraph_format.space_after = Pt(6)

    doc.add_paragraph()

    # ==========================================
    # 6. 列表样式设置
    # ==========================================
    doc.add_heading('三、数据集准备状态', level=1)
    items = [
        '原始数据：train.csv（3,000,888 行）已下载',
        '裁剪数据：train_subset.csv 已生成（277,860 行，5家门店 × 33品类 × 1684天）',
        '节假日数据：holidays_events.csv 已就位',
        '门店信息：stores.csv 已就位',
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('四、前端 Bug 修复记录', level=1)
    bugs = [
        ('analysis.js', 'renderTrendChart 数据字段错误', '修复图表横纵坐标字段映射'),
        ('prediction.js', '模型名大小写错误', 'arima/prophet/lstm → ARIMA/Prophet/LSTM'),
        ('report.js', '数据路径与置信区间字段名错误', '修改为 data.family 与 lower_ci/upper_ci'),
    ]

    table2 = doc.add_table(rows=1, cols=3)
    table2.style = 'Light Shading Accent 1'
    hdr2 = table2.rows[0].cells
    for i, text in enumerate(['文件', '问题描述', '修复方案']):
        hdr2[i].text = text
        hdr2[i].paragraphs[0].runs[0].font.bold = True

    for b in bugs:
        row = table2.add_row().cells
        for i, val in enumerate(b):
            row[i].text = val
            row[i].paragraphs[0].paragraph_format.space_before = Pt(6)
            row[i].paragraphs[0].paragraph_format.space_after = Pt(6)
            
    doc.add_paragraph()

    doc.add_heading('五、下一步计划（推进至 100%）', level=1)
    next_steps = [
        '补充测试数据文件（valid_gbk.csv、valid_xlsx.xlsx、short_series.csv 等）',
        '完善 tests/conftest.py 和各测试用例',
        '安装 Prophet 和 PyTorch 依赖，进行端到端联调测试',
        '运行 pytest tests/ -v 确保全部通过',
        '最终验收：上传数据 → 分析页图表 → 预测 BEVERAGES 品类 → 报告页展示',
    ]
    for step in next_steps:
        doc.add_paragraph(step, style='List Number')

    # 保存文档
    output_path = 'D:\\lml\\项目进度报告.docx'
    doc.save(output_path)
    print(f'商务美化版文档已生成：{output_path}')

if __name__ == '__main__':
    generate_report()