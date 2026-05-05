from __future__ import annotations

import argparse
import json
import re
import zipfile
from collections import Counter
from pathlib import Path
from xml.etree import ElementTree as ET

from docx import Document
from docx.enum.section import WD_SECTION


W_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
WP_NS = "{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}"


def pt(value):
    return None if value is None else round(value.pt, 2)


def cm(value):
    return None if value is None else round(value.cm, 2)


def para_text(paragraph):
    return paragraph.text.strip()


def para_style(paragraph):
    style = paragraph.style
    font = style.font if style is not None else None
    pf = style.paragraph_format if style is not None else None
    direct = paragraph.paragraph_format
    runs = paragraph.runs
    run_fonts = []
    run_sizes = []
    for run in runs:
        if run.text.strip():
            run_fonts.append(run.font.name)
            run_sizes.append(pt(run.font.size))
    return {
        "style": style.name if style is not None else None,
        "style_font": font.name if font is not None else None,
        "style_size_pt": pt(font.size) if font is not None else None,
        "style_bold": font.bold if font is not None else None,
        "style_line_spacing": pf.line_spacing if pf is not None else None,
        "direct_alignment": str(direct.alignment) if direct.alignment is not None else None,
        "direct_line_spacing": direct.line_spacing,
        "direct_space_before_pt": pt(direct.space_before),
        "direct_space_after_pt": pt(direct.space_after),
        "run_fonts": Counter(run_fonts).most_common(),
        "run_sizes_pt": Counter(run_sizes).most_common(),
    }


def extract_docx_xml_counts(path: Path):
    counts = {}
    with zipfile.ZipFile(path) as zf:
        names = set(zf.namelist())
        for name in [
            "word/document.xml",
            "word/settings.xml",
            "word/styles.xml",
            "word/header1.xml",
            "word/footer1.xml",
        ]:
            counts[name] = name in names
        root = ET.fromstring(zf.read("word/document.xml"))
    counts["page_breaks"] = len(root.findall(f".//{W_NS}br[@{W_NS}type='page']"))
    counts["section_breaks"] = len(root.findall(f".//{W_NS}sectPr"))
    counts["drawings"] = len(root.findall(f".//{W_NS}drawing"))
    counts["inline_drawings"] = len(root.findall(f".//{WP_NS}inline"))
    counts["tables"] = len(root.findall(f".//{W_NS}tbl"))
    counts["num_references"] = len(root.findall(f".//{W_NS}numPr"))
    counts["fld_characters"] = len(root.findall(f".//{W_NS}fldChar"))
    return counts


def summarize_docx(path: Path):
    doc = Document(path)
    sections = []
    for i, section in enumerate(doc.sections, 1):
        sections.append(
            {
                "index": i,
                "start_type": str(section.start_type),
                "page_width_cm": cm(section.page_width),
                "page_height_cm": cm(section.page_height),
                "top_margin_cm": cm(section.top_margin),
                "bottom_margin_cm": cm(section.bottom_margin),
                "left_margin_cm": cm(section.left_margin),
                "right_margin_cm": cm(section.right_margin),
                "header_distance_cm": cm(section.header_distance),
                "footer_distance_cm": cm(section.footer_distance),
                "different_first_page_header_footer": section.different_first_page_header_footer,
            }
        )

    paragraphs = [p for p in doc.paragraphs if para_text(p)]
    first_paragraphs = []
    for idx, p in enumerate(paragraphs[:80], 1):
        first_paragraphs.append(
            {
                "n": idx,
                "text": para_text(p)[:160],
                **para_style(p),
            }
        )

    heading_like = []
    patterns = [
        r"^摘\s*要$",
        r"^Abstract$",
        r"^目\s*录$",
        r"^第[一二三四五六七八九十]+章",
        r"^\d+(?:\.\d+)*\s+",
        r"^参考文献$",
        r"^致\s*谢$",
        r"^附\s*录",
    ]
    for idx, p in enumerate(doc.paragraphs, 1):
        text = para_text(p)
        if not text:
            continue
        if any(re.search(pattern, text, re.I) for pattern in patterns) or "标题" in (p.style.name if p.style else ""):
            heading_like.append(
                {
                    "paragraph_index": idx,
                    "text": text[:200],
                    **para_style(p),
                }
            )

    table_summary = []
    for ti, table in enumerate(doc.tables, 1):
        sample = []
        for row in table.rows[:3]:
            sample.append([cell.text.strip().replace("\n", " ")[:80] for cell in row.cells[:5]])
        table_summary.append(
            {
                "index": ti,
                "rows": len(table.rows),
                "cols": len(table.columns),
                "sample": sample,
            }
        )

    text = "\n".join(p.text for p in doc.paragraphs)
    issues = {
        "contains_placeholder": bool(re.search(r"(TODO|待补|此处|请补|xxx|XXX|模板|占位)", text)),
        "contains_question_marks_runs": "???" in text,
        "reference_count_estimate": len(re.findall(r"^\s*\[\d+\]", text, re.M)),
        "figure_caption_count": len(re.findall(r"图\s*\d+(?:[-.]\d+)?", text)),
        "table_caption_count": len(re.findall(r"表\s*\d+(?:[-.]\d+)?", text)),
        "keyword_line": next((line for line in text.splitlines() if "关键词" in line[:12]), ""),
    }

    return {
        "path": str(path),
        "sections": sections,
        "xml_counts": extract_docx_xml_counts(path),
        "paragraph_count": len(doc.paragraphs),
        "nonempty_paragraph_count": len(paragraphs),
        "first_paragraphs": first_paragraphs,
        "heading_like": heading_like[:120],
        "table_summary": table_summary,
        "quick_issues": issues,
    }


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("docx", type=Path)
    parser.add_argument("--out", type=Path)
    args = parser.parse_args()
    summary = summarize_docx(args.docx)
    payload = json.dumps(summary, ensure_ascii=False, indent=2)
    if args.out:
        args.out.parent.mkdir(parents=True, exist_ok=True)
        args.out.write_text(payload, encoding="utf-8")
    else:
        print(payload)


if __name__ == "__main__":
    main()
