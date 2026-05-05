from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("docx", type=Path)
    parser.add_argument("--out", type=Path, required=True)
    args = parser.parse_args()

    doc = Document(args.docx)
    lines = []
    for i, paragraph in enumerate(doc.paragraphs, 1):
        text = paragraph.text.strip()
        if text:
            lines.append(f"{i:04d}\t[{paragraph.style.name}]\t{text}")
    args.out.parent.mkdir(parents=True, exist_ok=True)
    args.out.write_text("\n".join(lines), encoding="utf-8")
    print(f"paragraphs={len(doc.paragraphs)} nonempty={len(lines)} out={args.out}")


if __name__ == "__main__":
    main()
