from __future__ import annotations

import re
import shutil
import zipfile
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
DRAFTS = ROOT / "drafts"
SOURCE_MD = DRAFTS / "毕业论文初稿.md"
OUTPUT_MD = DRAFTS / "毕业论文初稿.md"
OUTPUT_DOCX = ROOT / "商贸公司销售数据分析与智能预测系统_论文初稿.docx"
BACKUP_MD = DRAFTS / f"毕业论文初稿_升级前备份_{datetime.now():%Y%m%d_%H%M%S}.md"

TITLE = "商贸公司销售数据分析与智能预测系统"
SCHOOL = "北京信息科技大学"
COLLEGE = "计算机学院"
MAJOR = "数据科学与大数据技术"
AUTHOR = "刘明亮"
CLASS_NO = "数据GT2401（2024077370）"
ADVISOR = "李莉"
DATE_TEXT = "2026 年 5 月"
LATEST_TEST_RESULT = "44 passed in 45.92s"

CHAPTER_NUMS = {
    "1": "一",
    "2": "二",
    "3": "三",
    "4": "四",
    "5": "五",
    "6": "六",
    "7": "七",
}


def upgrade_markdown(source: str) -> str:
    """Refresh volatile evidence while preserving the manually drafted thesis."""
    if "## 摘要" not in source or "## 第 1 章" not in source:
        raise RuntimeError("Cannot find the expected thesis sections in source markdown")

    text = source
    text = re.sub(r"44 passed in \d+\.\d+s", LATEST_TEST_RESULT, text)
    text = re.sub(r"第 ([1-7]) 章", lambda m: f"第 {m.group(1)} 章", text)
    return text


def set_run_font(run, size: float = 10.5, bold: bool = False, font: str = "宋体") -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.bold = bold


def set_paragraph_format(paragraph, first_line: bool = True, line_spacing: float = 18) -> None:
    fmt = paragraph.paragraph_format
    fmt.line_spacing = Pt(line_spacing)
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    if first_line:
        fmt.first_line_indent = Cm(0.74)


def configure_section(section) -> None:
    section.page_height = Cm(29.7)
    section.page_width = Cm(21)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)
    section.header_distance = Cm(1.5)
    section.footer_distance = Cm(1.75)


def clear_header_footer(section) -> None:
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    section.header.paragraphs[0].text = ""
    section.footer.paragraphs[0].text = ""


def set_page_number_format(section, fmt: str, start: int = 1) -> None:
    sect_pr = section._sectPr
    pg_num = sect_pr.find(qn("w:pgNumType"))
    if pg_num is None:
        pg_num = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num)
    pg_num.set(qn("w:start"), str(start))
    pg_num.set(qn("w:fmt"), fmt)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_text = OxmlElement("w:t")
    fld_text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_sep, fld_text, fld_end])
    set_run_font(run, 9)


def set_header_footer(section, with_header: bool, page_fmt: str, start: int | None = None) -> None:
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    if start is not None:
        set_page_number_format(section, page_fmt, start)

    header = section.header.paragraphs[0]
    header.text = TITLE if with_header else ""
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in header.runs:
        set_run_font(run, 9, True)

    footer = section.footer.paragraphs[0]
    footer.text = ""
    add_page_number(footer)


def add_centered(doc, text: str, size: float = 10.5, bold: bool = False, font: str = "宋体") -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_format(p, first_line=False)
    run = p.add_run(text)
    set_run_font(run, size, bold, font)


def add_normal(doc, text: str, first_line: bool = True, size: float = 10.5, bold: bool = False) -> None:
    p = doc.add_paragraph()
    set_paragraph_format(p, first_line=first_line)
    run = p.add_run(text)
    set_run_font(run, size, bold)


def add_heading(doc, text: str, level: int) -> None:
    if level == 1:
        p = doc.add_paragraph(style="Heading 1")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_format(p, first_line=False)
        run = p.add_run(text)
        set_run_font(run, 16, True)
        return
    if level == 2:
        p = doc.add_paragraph(style="Heading 2")
        set_paragraph_format(p, first_line=False)
        run = p.add_run(text)
        set_run_font(run, 12, True)
        return
    p = doc.add_paragraph(style="Heading 3")
    set_paragraph_format(p, first_line=False)
    run = p.add_run(text)
    set_run_font(run, 10.5, True)


def add_toc(paragraph) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r'TOC \o "1-2" \h \z \u'
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_text = OxmlElement("w:t")
    fld_text.text = "请在 Word 中右键更新域以生成目录"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_sep, fld_text, fld_end])
    set_run_font(run)


def normalize_heading(text: str) -> str:
    match = re.match(r"第\s*([1-7])\s*章\s*(.+)", text)
    if match:
        return f"第{CHAPTER_NUMS[match.group(1)]}章 {match.group(2)}"
    return text


def clean_inline(text: str) -> str:
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"\*(.*?)\*", r"\1", text)
    return text.replace("`", "")


def add_markdown_table(doc, lines: list[str]) -> None:
    rows: list[list[str]] = []
    for line in lines:
        cells = [clean_inline(cell.strip()) for cell in line.strip().strip("|").split("|")]
        if all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells):
            continue
        rows.append(cells)
    if not rows:
        return

    width = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=width)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        for c_idx in range(width):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            text = row[c_idx] if c_idx < len(row) else ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if r_idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
            set_paragraph_format(p, first_line=False)
            run = p.add_run(text)
            set_run_font(run, 10.5, r_idx == 0)


def add_cover(doc: Document) -> None:
    add_centered(doc, SCHOOL, 26, False, "楷体")
    for _ in range(3):
        doc.add_paragraph()
    add_centered(doc, "毕 业 设 计（论 文）", 28, True, "楷体")
    for _ in range(4):
        doc.add_paragraph()
    for label, value in [
        ("题    目：", TITLE),
        ("学    院：", COLLEGE),
        ("专    业：", MAJOR),
        ("学生姓名：", f"{AUTHOR}    班级/学号：{CLASS_NO}"),
        ("指导老师/督导老师：", ADVISOR),
        ("完成时间：", DATE_TEXT),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_format(p, first_line=False, line_spacing=24)
        run = p.add_run(f"{label}{value}")
        set_run_font(run, 14)


def build_docx(md: str, path: Path) -> None:
    doc = Document()
    configure_section(doc.sections[0])
    clear_header_footer(doc.sections[0])

    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    styles["Normal"].font.size = Pt(10.5)

    add_cover(doc)
    doc.add_page_break()

    start = md.find("## 原创性声明和版权授权声明")
    if start == -1:
        raise RuntimeError("Cannot find declaration section")
    lines = md[start:].splitlines()

    in_code = False
    body_started = False
    skip_manual_toc = False
    table_buffer: list[str] = []

    def flush_table() -> None:
        nonlocal table_buffer
        if table_buffer:
            add_markdown_table(doc, table_buffer)
            table_buffer = []

    for raw in lines:
        line = raw.rstrip()
        if not line:
            flush_table()
            continue

        if line.startswith("> "):
            continue

        if line.strip() == "---":
            flush_table()
            continue

        if skip_manual_toc and not line.startswith("## "):
            continue

        if line.startswith("```"):
            flush_table()
            in_code = not in_code
            continue

        if in_code:
            p = doc.add_paragraph()
            set_paragraph_format(p, first_line=False)
            run = p.add_run(line)
            set_run_font(run, 9, False, "宋体")
            continue

        if line.startswith("|") and "|" in line[1:]:
            table_buffer.append(line)
            continue
        flush_table()

        if line.startswith("## "):
            text = normalize_heading(line[3:].strip())
            if text == "封面":
                continue
            if text == "原创性声明和版权授权声明":
                add_centered(doc, text, 16, True)
                continue
            if text in {"摘要", "Abstract"}:
                if text == "摘要":
                    section = doc.add_section(WD_SECTION.NEW_PAGE)
                    configure_section(section)
                    set_header_footer(section, with_header=False, page_fmt="upperRoman", start=1)
                else:
                    doc.add_page_break()
                add_heading(doc, text, 1)
                continue
            if text == "目录":
                doc.add_page_break()
                add_heading(doc, "目  录", 1)
                p = doc.add_paragraph()
                add_toc(p)
                skip_manual_toc = True
                continue
            if text.startswith("第一章") and not body_started:
                skip_manual_toc = False
                body_started = True
                section = doc.add_section(WD_SECTION.NEW_PAGE)
                configure_section(section)
                set_header_footer(section, with_header=True, page_fmt="decimal", start=1)
                add_heading(doc, text, 1)
                continue
            if body_started:
                doc.add_page_break()
            add_heading(doc, text, 1)
            continue

        if line.startswith("### "):
            add_heading(doc, normalize_heading(line[4:].strip()), 2)
            continue

        if line.startswith("#### "):
            add_heading(doc, line[5:].strip(), 3)
            continue

        text = clean_inline(line)
        if re.match(r"^\[\d+\]", text):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.74)
            p.paragraph_format.first_line_indent = Cm(-0.74)
            p.paragraph_format.line_spacing = Pt(18)
            run = p.add_run(text)
            set_run_font(run)
        elif re.match(r"^\d+\.\s+", text) or text.startswith("- "):
            p = doc.add_paragraph()
            set_paragraph_format(p, first_line=True)
            run = p.add_run(text)
            set_run_font(run)
        else:
            add_normal(doc, text)

    flush_table()
    doc.save(path)
    enable_update_fields(path)


def enable_update_fields(docx_path: Path) -> None:
    settings_name = "word/settings.xml"
    tmp_path = docx_path.with_suffix(".tmp.docx")
    with zipfile.ZipFile(docx_path, "r") as src, zipfile.ZipFile(tmp_path, "w", zipfile.ZIP_DEFLATED) as dst:
        for item in src.infolist():
            data = src.read(item.filename)
            if item.filename == settings_name:
                xml = data.decode("utf-8")
                if "w:updateFields" not in xml:
                    xml = xml.replace("</w:settings>", '<w:updateFields w:val="true"/></w:settings>')
                data = xml.encode("utf-8")
            elif item.filename.startswith("word/") and item.filename.endswith(".xml"):
                xml = data.decode("utf-8")
                xml = re.sub(r'<w:color\b[^>]*/>', '<w:color w:val="000000"/>', xml)
                xml = re.sub(r'<w:color\b[^>]*>.*?</w:color>', '<w:color w:val="000000"/>', xml)
                data = xml.encode("utf-8")
            dst.writestr(item, data)
    tmp_path.replace(docx_path)


def force_black_text(docx_path: Path) -> None:
    tmp_path = docx_path.with_suffix(".black.tmp.docx")
    with zipfile.ZipFile(docx_path, "r") as src, zipfile.ZipFile(tmp_path, "w", zipfile.ZIP_DEFLATED) as dst:
        for item in src.infolist():
            data = src.read(item.filename)
            if item.filename.startswith("word/") and item.filename.endswith(".xml"):
                xml = data.decode("utf-8")
                xml = re.sub(r'<w:color\b[^>]*/>', '<w:color w:val="000000"/>', xml)
                xml = re.sub(r'<w:color\b[^>]*>.*?</w:color>', '<w:color w:val="000000"/>', xml)
                data = xml.encode("utf-8")
            dst.writestr(item, data)
    tmp_path.replace(docx_path)


def main() -> None:
    source = SOURCE_MD.read_text(encoding="utf-8")
    existing_backups = list(DRAFTS.glob("毕业论文初稿_升级前备份_*.md"))
    if not existing_backups:
        shutil.copy2(SOURCE_MD, BACKUP_MD)

    upgraded = upgrade_markdown(source)
    OUTPUT_MD.write_text(upgraded, encoding="utf-8", newline="\n")
    build_docx(upgraded, OUTPUT_DOCX)
    print(f"markdown={OUTPUT_MD}")
    print(f"docx={OUTPUT_DOCX}")
    print(f"chars={len(upgraded)}")


if __name__ == "__main__":
    main()
